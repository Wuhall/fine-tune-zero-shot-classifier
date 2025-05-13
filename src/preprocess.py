import os
from docx import Document
import openpyxl

def is_empty_cell(cell):
    return not any(paragraph.text.strip() for paragraph in cell.paragraphs)

def get_unique_content(cells):
    # 获取非空内容并去重
    contents = [cell.text.strip() for cell in cells if not is_empty_cell(cell)]
    return list(dict.fromkeys(contents))  # 保持顺序去重

def analyze_tables(docx_path, workbook, sheet):
    print(f"\n分析文件: {os.path.basename(docx_path)}")
    print("=" * 100)
    
    doc = Document(docx_path)
    all_rows_content = []  # 存储所有行的内容
    
    for table_index, table in enumerate(doc.tables, 1):
        # 检查表格是否为空
        has_content = False
        for row in table.rows:
            if any(not is_empty_cell(cell) for cell in row.cells):
                has_content = True
                break
        
        if not has_content:
            continue
        
        for row_index, row in enumerate(table.rows, 1):
            # 获取去重后的内容
            unique_contents = get_unique_content(row.cells)
            if not unique_contents:  # 如果行是空的，跳过
                continue
            
            # 将内容添加到列表中
            content_str = " ".join(unique_contents)
            if len(content_str) < 50:
                continue
            all_rows_content.append((content_str, row_index, table_index))
    
    # 将内容写入 Excel 文件
    for content, _, _ in all_rows_content:
        sheet.append([os.path.basename(docx_path), content])

# 创建 Excel 工作簿
output_path = "/Users/lucas/Desktop/FileExtractorData.xlsx"
workbook = openpyxl.Workbook()
sheet = workbook.active
sheet.title = "ExtractedData"
sheet.append(["file_name", "paragraph"])  # 添加表头

# 指定目录路径
directory = "/Users/lucas/Downloads/test"

# 遍历目录下的所有文件
for filename in os.listdir(directory):
    # 忽略.~开头的文件和隐藏文件
    if filename.endswith(".docx") and not filename.startswith(".~") and not filename.startswith("."):
        file_path = os.path.join(directory, filename)
        try:
            analyze_tables(file_path, workbook, sheet)
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")

# 保存 Excel 文件
workbook.save(output_path)
print(f"\n数据已保存到 {output_path}")